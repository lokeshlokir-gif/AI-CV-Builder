#!/usr/bin/env python3
"""AI CV Builder - Streamlit Web App
Region-aware CV (14 regions), Word/PDF export, clickable home cards,
130+ mock interview domains, experience-level tailoring, anti-repeat AI
with persistent history, mobile-safe sidebar, emoji-stripped exports,
Gemini + OpenAI + Claude providers, open My Library, voice input + AI polish."""
import streamlit as st
import streamlit.components.v1 as components
import json, re, time, io, ssl, random, uuid
import urllib.request, urllib.error
from datetime import datetime

# ============================================================
# SSL SETUP
# ============================================================
def _get_ssl():
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where()), "certifi"
    except: pass
    try:
        ctx = ssl.create_default_context()
        urllib.request.urlopen("https://www.google.com", timeout=5, context=ctx)
        return ctx, "system"
    except: pass
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    return ctx, "unverified"
SSL_CTX, SSL_M = _get_ssl()

# ============================================================
# OPTIONAL LIBRARIES
# ============================================================
HAS_PDF = False
try:
    from pypdf import PdfReader
    HAS_PDF = True
except:
    try:
        from PyPDF2 import PdfReader
        HAS_PDF = True
    except: pass

HAS_DOCX = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except: pass

HAS_RL = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    HAS_RL = True
except: pass

HAS_REQ = False
try:
    import requests as rq
    HAS_REQ = True
except: pass

# ============================================================
# MODELS — Gemini + OpenAI + Claude
# ============================================================
GEMINI_MODELS = [
    ("gemini-2.5-flash", "2.5 Flash (FREE)"),
    ("gemini-2.5-flash-lite", "2.5 Flash-Lite"),
    ("gemini-2.5-pro", "2.5 Pro"),
    ("gemini-2.0-flash", "2.0 Flash"),
]
OPENAI_MODELS = [
    ("gpt-4o", "GPT-4o (Best)"),
    ("gpt-4o-mini", "GPT-4o Mini"),
    ("gpt-4-turbo", "GPT-4 Turbo"),
    ("gpt-3.5-turbo", "GPT-3.5 Turbo"),
]
CLAUDE_MODELS = [
    ("claude-opus-4-20250514", "Claude Opus 4 (Best)"),
    ("claude-sonnet-4-20250514", "Claude Sonnet 4"),
    ("claude-3-5-sonnet-20241022", "Claude 3.5 Sonnet"),
    ("claude-3-5-haiku-20241022", "Claude 3.5 Haiku (Fast)"),
]
FB_ = [m[0] for m in GEMINI_MODELS]
MR = 3
RDL = 8

# ============================================================
# MOCK INTERVIEW DOMAINS — 130+ in categories
# ============================================================
DOMAIN_GROUPS = {
    "General": ["General"],
    "Technology": [
        "AI / Machine Learning", "Data Science / Analytics", "Cyber Security",
        "Cloud / DevOps", "Software Engineering", "Web Development",
        "Mobile Development", "Database / SQL", "Networking",
        "Blockchain / Web3", "AR / VR", "Embedded Systems",
    ],
    "Management & Business": [
        "Project Management", "Agile / Scrum", "Product Management",
        "Business Analyst", "Operations Management", "Strategy Consulting",
    ],
    "Finance & Sales": [
        "Finance / Banking", "Investment / Trading", "Accounting / CA",
        "Sales / Business Dev", "Marketing / Digital", "HR / Recruitment",
        "Supply Chain / Logistics", "Insurance", "Wealth Management",
    ],
    "Engineering": [
        "Computer Science Engineering (CSE)",
        "Information Science Engineering (ISE)",
        "Information Technology (IT)",
        "Electronics & Communication (ECE)",
        "Electrical & Electronics (EEE)",
        "Mechanical Engineering",
        "Civil Engineering",
        "Chemical Engineering",
        "Aerospace Engineering",
        "Automobile Engineering",
        "Industrial Engineering",
        "Biomedical Engineering",
        "Robotics & Mechatronics",
        "Mining Engineering",
        "Petroleum Engineering",
        "Marine Engineering",
        "Textile Engineering",
        "Metallurgy & Materials",
        "Biotechnology / Genetic Eng",
        "Architecture",
    ],
    "Medical & Healthcare": [
        "MBBS / Medicine", "BDS / Dentistry", "BPT / Physiotherapy",
        "Nursing (B.Sc)", "Pharma-B / Pharmacy", "Pharma-D",
        "Veterinary", "Public Health", "Medical Lab Technology",
        "Radiology", "Optometry",
        "AYUSH - BAMS (Ayurveda)", "AYUSH - BHMS (Homeopathy)",
        "AYUSH - BUMS (Unani)", "Yoga & Naturopathy",
        "Nutrition & Dietetics", "Occupational Therapy",
    ],
    "Education & Teaching": [
        "Teaching / School", "Higher Education / Lecturer",
        "Special Education", "Education Administration", "Tutoring",
        "B.Ed / D.Ed", "Early Childhood / Montessori",
    ],
    "Indian Education Streams": [
        "PUC PCMB", "PUC PCMC", "PUC Commerce", "PUC Arts",
        "Diploma / Polytechnic", "ITI",
        "BCA", "BBA / BBM", "B.A", "B.Com", "B.Sc",
        "B.Tech / B.E", "MBA", "M.Sc", "M.A", "M.Com",
        "M.Tech / M.E", "Ph.D / Research",
    ],
    "Law & Legal": [
        "Law / Legal (General)", "Corporate Law", "Criminal Law",
        "Civil Law", "IP Law", "Tax Law", "Constitutional Law",
        "Cyber Law",
    ],
    "Agriculture & Environment": [
        "Agriculture / Agronomy", "Horticulture", "Forestry",
        "Environmental Science", "Fisheries", "Dairy Technology",
        "Food Technology", "Sericulture",
    ],
    "Arts & Social Sciences": [
        "Geography", "History", "Psychology", "Sociology",
        "Political Science", "Economics", "Anthropology",
        "Philosophy", "Literature", "Linguistics",
        "Social Work (MSW)",
    ],
    "Creative & Media": [
        "Journalism", "Graphic Design", "UI/UX Design",
        "Photography", "Film Production", "Music",
        "Fashion Design", "Interior Design",
        "Animation", "VFX", "Game Design",
        "Content Creation / YouTube", "Copywriting",
    ],
    "Services & Hospitality": [
        "Hotel Management", "Travel & Tourism",
        "Aviation / Air Hostess", "Event Management",
        "Real Estate", "Retail Management", "Customer Service",
    ],
    "Government & Civil Services": [
        "UPSC / Civil Services", "State PSC",
        "Banking Exams (IBPS/SBI)", "SSC (CGL/CHSL)",
        "Railway Exams (RRB)", "Defence (NDA/CDS)",
        "Police / Constable",
    ],
    "Trade & Skilled": [
        "Plumbing / Electrical Trade", "Auto Mechanic",
        "Carpentry", "Welding", "Hair Stylist / Beautician",
        "Tailoring / Stitching", "Painting / Construction",
        "Chef / Cooking", "Driving / Heavy Vehicle",
    ],
}

EXP_LEVELS = [
    "🎓 Student / Fresher (0 yrs)",
    "👶 Junior (0-2 yrs)",
    "👨‍💻 Mid-level (2-5 yrs)",
    "🧑‍🔧 Senior (5-10 yrs)",
    "👑 Lead / Principal (10+ yrs)",
]

# ============================================================
# REGION DATA (14 regions)
# ============================================================
RD_ = {}
RD_['🇬🇧 United Kingdom'] = dict(code='UK', term='CV', pages='2 pages A4', photo='NO photo - never include photo or DOB', pn=False, pi='Name, phone, email, LinkedIn (NO photo, NO DOB, NO nationality)', sec='Personal Profile, Key Skills, Work Experience, Education, Additional', sty='Achievement-focused with measurable results, UK English', sp='Personal Profile (3-4 lines) at top is ESSENTIAL.', ats='ATS-critical.', av='Photos, DOB, nationality, references', cp=(0, 51, 102))
RD_['🇺🇸 United States'] = dict(code='US', term='Resume', pages='1-2 pages Letter', photo='NO photo - never include photo or age', pn=False, pi='Name, phone, email, LinkedIn', sec='Summary, Core Competencies, Professional Experience, Education, Certifications', sty='STAR method with quantifiable results, US English', sp='Quantify everything with %, $, numbers.', ats='ATS-CRITICAL.', av='Photos, age, DOB, SSN', cp=(0, 51, 102))
RD_['🇨🇦 Canada'] = dict(code='CA', term='Resume', pages='1-2 pages', photo='NO photo', pn=False, pi='Name, phone, email, LinkedIn, city/province', sec='Summary, Qualifications, Experience, Education, Volunteer Experience', sty='Achievement-focused, bilingual EN/FR a bonus', sp='Volunteer experience is valued.', ats='ATS required.', av='Photos, DOB, SIN', cp=(139, 0, 0))
RD_['🇩🇪 Germany'] = dict(code='DE', term='Lebenslauf', pages='2-3 pages A4', photo='YES - include professional headshot photo placeholder at top-right (3.5x4.5cm)', pn=True, pi='Name, address, phone, email, DOB, nationality, marital status, photo placeholder', sec='Personliche Daten, Berufserfahrung, Ausbildung, Kenntnisse, Sprachen', sty='Formal, strict reverse-chronological, formal German tone', sp='Photo at top-right is EXPECTED. Include personal data (DOB, nationality).', ats='Growing ATS.', av='Casual tone, gaps', cp=(15, 52, 96))
RD_['🇫🇷 France'] = dict(code='FR', term='CV', pages='1-2 pages A4', photo='OPTIONAL - photo is common', pn=True, pi='Name, address, phone, email, age/DOB, optional photo', sec='Etat Civil, Experience Professionnelle, Formation, Competences, Langues (CEFR)', sty='Formal French CV with CEFR levels', sp='CEFR language levels mandatory.', ats='Growing ATS.', av='Overly casual', cp=(0, 35, 102))
RD_['🇪🇸 Spain'] = dict(code='ES', term='CV', pages='1-2 pages A4', photo='EXPECTED - small professional photo at top', pn=True, pi='Name, address, phone, email, DNI/NIE, DOB, photo placeholder', sec='Datos Personales, Experiencia, Formacion, Idiomas, Competencias', sty='Europass-style common', sp='Photo expected.', ats='Europass.', av='Missing personal data', cp=(170, 0, 0))
RD_['🇮🇹 Italy'] = dict(code='IT', term='CV', pages='2-3 pages A4', photo='EXPECTED - professional photo at top-right', pn=True, pi='Name, address, phone, email, DOB, nationality, photo placeholder', sec='Dati Personali, Esperienza, Istruzione, Competenze, Lingue', sty='Europass format with GDPR clause', sp='Photo expected. GDPR privacy clause MANDATORY at end.', ats='Europass.', av='Missing GDPR clause', cp=(0, 102, 51))
RD_['🇮🇳 India'] = dict(code='IN', term='Resume/CV', pages='2-3 pages A4', photo='OPTIONAL', pn=False, pi='Name, phone, email, LinkedIn, city', sec='Career Objective, Skills, Experience, Education (detailed), Projects, Declaration', sty='Education-focused, tech skills prominent', sp='Declaration statement at bottom is traditional. Education detailed.', ats='Growing ATS.', av='Missing declaration', cp=(255, 103, 0))
RD_['🇦🇪 UAE / Gulf'] = dict(code='AE', term='CV', pages='2-4 pages A4', photo='RECOMMENDED - professional photo at top', pn=True, pi='Name, phone, email, NATIONALITY (mandatory), VISA STATUS (mandatory), DOB, photo', sec='Personal Details, Career Objective, Summary, Experience, Education, Skills, Languages', sty='Detailed CV, nationality and visa status critical', sp='Nationality and visa status MANDATORY at top.', ats='Less ATS.', av='Missing nationality/visa', cp=(0, 100, 0))
RD_['🇯🇵 Japan'] = dict(code='JP', term='Rirekisho', pages='1-2 pages STRICT format', photo='REQUIRED - professional photo at top-right (3x4cm) in dedicated box', pn=True, pi='Name (kanji+furigana), address, phone, email, DOB, gender, photo', sec='Personal Info (with PHOTO BOX), Education (chronological), Work History (chronological), Qualifications, Motivation', sty='STRICT Rirekisho - chronological (oldest first), photo top-right, seal at bottom', sp='Photo REQUIRED at top-right. Chronological order (oldest first) MANDATORY. Mark a PHOTO BOX area clearly.', ats='Strict template.', av='Reverse-chronological order', cp=(153, 0, 0))
RD_['🇨🇳 China'] = dict(code='CN', term='Resume', pages='1 page A4', photo='EXPECTED - professional photo at top-right', pn=True, pi='Name, phone, email, DOB, gender, hukou, photo', sec='Personal Information, Education, Work Experience, Skills, Awards, Self-evaluation', sty='Concise one-pager, highlight 985/211 universities', sp='Photo expected. One page strictly.', ats='Growing.', av='Multi-page', cp=(204, 0, 0))
RD_['🇦🇺 Australia'] = dict(code='AU', term='Resume/CV', pages='2-3 pages A4', photo='NO photo', pn=False, pi='Name, phone, email, LinkedIn, city', sec='Career Summary, Key Skills, Employment History, Education, Referees', sty='Achievement-focused, named referees expected', sp='Referees with names and contact details EXPECTED at end.', ats='ATS recommended.', av='Photos, missing referees', cp=(0, 0, 128))
RD_['🇸🇬 Singapore'] = dict(code='SG', term='Resume/CV', pages='1-2 pages A4', photo='OPTIONAL', pn=False, pi='Name, phone, email, nationality, work pass status', sec='Summary, Skills, Experience, Education, Certifications, Languages', sty='Concise, work pass status important', sp='Work pass status (EP/SP/PR) is IMPORTANT.', ats='ATS widely used.', av='Missing work pass info', cp=(204, 0, 0))
RD_['🇪🇺 EU (Europass)'] = dict(code='EU', term='Europass CV', pages='2-3 pages A4', photo='OPTIONAL', pn=False, pi='Name, address, phone, email, nationality', sec='Personal Information, Work Experience, Education, Language Skills (CEFR), Digital Skills', sty='Standardised Europass format', sp='CEFR language levels required.', ats='Machine-readable.', av='Non-standard', cp=(0, 51, 153))
RL = list(RD_.keys())

# ============================================================
# STRIP EMOJIS / SPECIAL UNICODE FOR PDF/WORD SAFETY
# ============================================================
def strip_emojis(text):
    if not text:
        return ""
    text = re.sub(r'[\U0001F000-\U0001FFFF]', '', text)
    text = re.sub(r'[\U0001F1E6-\U0001F1FF]', '', text)
    text = re.sub(r'[\u2500-\u27BF]', '', text)
    text = re.sub(r'[\u2B00-\u2BFF]', '', text)
    text = re.sub(r'[\u2300-\u23FF]', '', text)
    text = re.sub(r'[\uFE00-\uFE0F]', '', text)
    text = text.replace('\u200d', '')
    for ch in '═━─━│┃┌┐└┘├┤┬┴┼╔╗╚╝╠╣╦╩╬':
        text = text.replace(ch, '')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

# ============================================================
# EXTRACT CANDIDATE NAME FOR FILENAMES
# ============================================================
def extract_name(text):
    if not text:
        return "Candidate"
    SK = ["SUMMARY", "SKILL", "EXPERIENCE", "EDUCATION", "CERTIF", "PROFILE",
          "OBJECTIVE", "PERSONAL", "CAREER", "LANGUAGES", "ADDITIONAL",
          "DECLARATION", "REFERENCES", "QUALIFICATIONS", "WORK HISTORY",
          "MOTIVATION", "DEAR", "TO WHOM", "HIRING MANAGER", "COVER LETTER",
          "SCORE", "ANALYSIS", "FEEDBACK", "QUESTION", "ANSWER", "STRENGTH",
          "WEAKNESS", "IDEAL", "TIPS", "INSIGHT", "ACTION", "RECOMMENDATION"]
    raw = [l for l in strip_emojis(text).split("\n") if l.strip()]
    candidates = raw[:10] + raw[-5:]
    for line in candidates:
        c = re.sub(r"#+\s*", "", line).strip()
        c = c.strip("*").strip("- •▪").strip()
        c = re.sub(r"^(sincerely|regards|best regards|yours sincerely|"
                   r"yours truly|kind regards|warm regards),?\s*", "",
                   c, flags=re.I).strip()
        if not c:
            continue
        u = c.upper().replace(":", "")
        if any(k in u for k in SK):
            continue
        if c.lower().startswith(("here", "below", "i have", "this is",
                                  "based on", "dear", "to whom")):
            continue
        if len(c) < 2 or len(c) > 60:
            continue
        if not re.match(r"^[A-Za-z][A-Za-z\s\.\-']{1,}$", c):
            continue
        safe = re.sub(r"[^A-Za-z0-9_\-]+", "_", c).strip("_")
        if safe:
            return safe
    return "Candidate"

# ============================================================
# NETWORK / AI CALLS — Gemini + OpenAI + Claude
# ============================================================
def _post(url, data, headers=None, to=120):
    b = json.dumps(data).encode()
    hdrs = {"Content-Type": "application/json"}
    if headers: hdrs.update(headers)
    r = urllib.request.Request(url, data=b, headers=hdrs, method="POST")
    try:
        with urllib.request.urlopen(r, timeout=to, context=SSL_CTX) as rp:
            return rp.status, rp.read().decode()
    except urllib.error.HTTPError as e:
        return e.code, e.read().decode()
    except Exception as e:
        return 0, str(e)

def _gem(pr, key, model, to=120, temperature=0.7):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
    pl = {"contents": [{"parts": [{"text": pr}]}],
          "generationConfig": {"temperature": temperature,
                               "maxOutputTokens": 8192,
                               "topP": 0.95}}
    if HAS_REQ:
        try:
            r = rq.post(url, json=pl, timeout=to, headers={"Content-Type": "application/json"})
            sc, bd = r.status_code, r.text
        except:
            try:
                r = rq.post(url, json=pl, timeout=to, headers={"Content-Type": "application/json"}, verify=False)
                sc, bd = r.status_code, r.text
            except:
                sc, bd = _post(url, pl, to=to)
    else:
        sc, bd = _post(url, pl, to=to)
    if sc == 200:
        res = json.loads(bd) if isinstance(bd, str) else bd
        return True, res["candidates"][0]["content"]["parts"][0]["text"]
    elif sc == 429: return False, f"RATE_LIMITED|{model}"
    elif sc == 404: return False, f"MODEL_NOT_FOUND|{model}"
    else: return False, f"[Error {sc}] {bd[:200]}"

def _openai(pr, key, model, to=120, temperature=0.7):
    url = "https://api.openai.com/v1/chat/completions"
    pl = {"model": model, "messages": [{"role": "user", "content": pr}],
          "temperature": temperature, "max_tokens": 4096}
    hdrs = {"Content-Type": "application/json", "Authorization": f"Bearer {key}"}
    if HAS_REQ:
        try:
            r = rq.post(url, json=pl, timeout=to, headers=hdrs)
            sc, bd = r.status_code, r.text
        except:
            sc, bd = _post(url, pl, headers=hdrs, to=to)
    else:
        sc, bd = _post(url, pl, headers=hdrs, to=to)
    if sc == 200:
        res = json.loads(bd) if isinstance(bd, str) else bd
        return True, res["choices"][0]["message"]["content"]
    elif sc == 401: return False, "Invalid OpenAI key."
    else: return False, f"[Error {sc}] {bd[:200]}"

def _claude(pr, key, model, to=120, temperature=0.7):
    """Anthropic Claude API call."""
    url = "https://api.anthropic.com/v1/messages"
    pl = {"model": model, "max_tokens": 4096,
          "temperature": min(temperature, 1.0),
          "messages": [{"role": "user", "content": pr}]}
    hdrs = {"Content-Type": "application/json",
            "x-api-key": key,
            "anthropic-version": "2023-06-01"}
    if HAS_REQ:
        try:
            r = rq.post(url, json=pl, timeout=to, headers=hdrs)
            sc, bd = r.status_code, r.text
        except:
            sc, bd = _post(url, pl, headers=hdrs, to=to)
    else:
        sc, bd = _post(url, pl, headers=hdrs, to=to)
    if sc == 200:
        res = json.loads(bd) if isinstance(bd, str) else bd
        try:
            return True, res["content"][0]["text"]
        except:
            return False, f"[Parse Error] {str(res)[:300]}"
    elif sc == 401: return False, "Invalid Claude/Anthropic key."
    else: return False, f"[Error {sc}] {bd[:200]}"

def ai_call(pr, key, prov="gemini", model="gemini-2.5-flash", temperature=0.7):
    try:
        if prov == "openai":
            ok, t = _openai(pr, key, model, temperature=temperature)
            return t
        if prov == "claude":
            ok, t = _claude(pr, key, model, temperature=temperature)
            return t
        for a in range(1, MR + 1):
            ok, t = _gem(pr, key, model, temperature=temperature)
            if ok: return t
            if "RATE_LIMITED" not in t and "MODEL_NOT_FOUND" not in t: return t
            if "MODEL_NOT_FOUND" in t: break
            if a < MR: time.sleep(RDL)
        for fb in FB_:
            if fb == model: continue
            ok, t = _gem(pr, key, fb, temperature=temperature)
            if ok: return f"[Used:{fb}]\n\n{t}"
            if "RATE_LIMITED" not in t: return t
            time.sleep(2)
        return "All models failed."
    except Exception as e:
        return f"[Error] {e}"

# ============================================================
# FILE READER
# ============================================================
def read_file(uploaded):
    if uploaded is None: return ""
    name = uploaded.name.lower()
    if name.endswith(".pdf"):
        if not HAS_PDF: return "Install pypdf"
        reader = PdfReader(io.BytesIO(uploaded.read()))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    elif name.endswith(".docx"):
        if not HAS_DOCX: return "Install python-docx"
        doc = Document(io.BytesIO(uploaded.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    elif name.endswith(".txt"):
        return uploaded.read().decode("utf-8", errors="ignore")
    return ""


# ============================================================
# WORD EXPORT
# ============================================================
def make_docx(text, region):
    if not HAS_DOCX: return None
    text = strip_emojis(text)
    ri = RD_.get(region, {})
    cp = ri.get("cp", (0, 51, 102))
    has_photo = ri.get("pn", False)
    PC = RGBColor(*cp)
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2);  s.right_margin = Cm(2)
    lines = [l for l in text.split("\n") if l.strip()]
    SK = ["SUMMARY", "SKILL", "EXPERIENCE", "EDUCATION", "CERTIF", "PROFILE",
          "OBJECTIVE", "PERSONAL", "CAREER", "LANGUAGES", "ADDITIONAL",
          "DECLARATION", "REFERENCES", "QUALIFICATIONS", "WORK HISTORY", "MOTIVATION"]
    name = ""; cs = 0
    for idx, line in enumerate(lines):
        c = re.sub(r"#+\s*", "", line.strip()).strip("*").strip()
        if not c: continue
        u = c.upper().replace(":", "")
        if not name and not any(k in u for k in SK) and len(c) > 1 and \
           not c.lower().startswith(("here", "below", "i have", "this is", "based on")):
            name = c; cs = idx + 1; break
    if has_photo and name:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        lc = table.rows[0].cells[0]; lc.width = Cm(12)
        p = lc.paragraphs[0]; r = p.add_run(name)
        r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = PC
        rc = table.rows[0].cells[1]; rc.width = Cm(4)
        pp = rc.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(20)
        pp.paragraph_format.space_after = Pt(20)
        pp.add_run("[ PHOTO ]\n3.5 x 4.5cm").font.size = Pt(10)
    else:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name); r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = PC
    for i in range(cs, len(lines)):
        line = lines[i].strip()
        if not line: continue
        if line.lower().startswith(("here is", "below is", "i have", "this is", "based on")):
            continue
        c = re.sub(r"#+\s*", "", line).strip()
        cfc = c.strip("*").strip()
        if not cfc: continue
        u = cfc.upper().replace(":", "")
        is_h = False; ht = cfc
        if line.startswith("#"):
            is_h = True
        elif line.startswith("**") and line.endswith("**") and len(line) > 4:
            ht = line.strip("*").strip()
            if any(k in ht.upper().replace(":", "") for k in SK): is_h = True
        elif len(cfc) >= 3 and not cfc.startswith(("-", "*")):
            al = [x for x in cfc if x.isalpha()]
            if al and sum(1 for x in al if x.isupper()) / len(al) > 0.7 and any(k in u for k in SK):
                is_h = True
        if is_h:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(ht.upper())
            r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = PC
        elif cfc.startswith(("-", "*")):
            bt = cfc.lstrip("-* ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.add_run("• ").font.size = Pt(10)
            r = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", bt))
            r.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.+?\*\*)", cfc):
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part.strip("*"))
                    r.font.size = Pt(10); r.font.bold = True
                else:
                    cl2 = re.sub(r"\*(.+?)\*", r"\1", part)
                    if cl2.strip():
                        r = p.add_run(cl2); r.font.size = Pt(10)
    if ri.get("code") == "IT":
        doc.add_paragraph()
        r = doc.add_paragraph().add_run(
            "Autorizzo il trattamento dei dati personali ai sensi del GDPR (UE 2016/679).")
        r.font.size = Pt(8); r.italic = True
    if ri.get("code") == "IN":
        doc.add_paragraph()
        r = doc.add_paragraph().add_run(
            "Declaration: I hereby declare that all information provided above is true to the best of my knowledge.")
        r.font.size = Pt(9); r.italic = True
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI CV Builder - " + datetime.now().strftime("%d %B %Y"))
    r.font.size = Pt(7); r.font.color.rgb = RGBColor(150, 150, 150); r.italic = True
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

# ============================================================
# PDF EXPORT
# ============================================================
def make_pdf(text, region):
    if not HAS_RL: return None
    text = strip_emojis(text)
    ri = RD_.get(region, {})
    cp = ri.get("cp", (0, 51, 102))
    has_photo = ri.get("pn", False)
    color_hex = "#%02X%02X%02X" % cp
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm,
                            leftMargin=2*cm, rightMargin=2*cm)
    ss = getSampleStyleSheet()
    ss.add(ParagraphStyle("CVName", fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=6, textColor=color_hex, fontName="Helvetica-Bold"))
    ss.add(ParagraphStyle("CVNameL", fontSize=20, leading=24, alignment=TA_LEFT, spaceAfter=6, textColor=color_hex, fontName="Helvetica-Bold"))
    ss.add(ParagraphStyle("CVHead", fontSize=12, leading=14, spaceBefore=12, spaceAfter=4, textColor=color_hex, fontName="Helvetica-Bold"))
    ss.add(ParagraphStyle("CVBody", fontSize=10, leading=13, spaceAfter=2, fontName="Helvetica"))
    ss.add(ParagraphStyle("CVBullet", fontSize=10, leading=13, leftIndent=20, spaceAfter=1, fontName="Helvetica"))
    ss.add(ParagraphStyle("CVPhoto", fontSize=9, leading=12, alignment=TA_CENTER, fontName="Helvetica"))
    lines = [l for l in text.split("\n") if l.strip()]
    SK = ["SUMMARY", "SKILL", "EXPERIENCE", "EDUCATION", "CERTIF", "PROFILE",
          "OBJECTIVE", "PERSONAL", "CAREER", "LANGUAGES", "ADDITIONAL",
          "DECLARATION", "REFERENCES", "QUALIFICATIONS", "WORK HISTORY", "MOTIVATION"]
    name = ""; cs = 0
    for idx, line in enumerate(lines):
        c = re.sub(r"#+\s*", "", line.strip()).strip("*").strip()
        if not c: continue
        u = c.upper().replace(":", "")
        if not name and not any(k in u for k in SK) and len(c) > 1 and \
           not c.lower().startswith(("here", "below", "i have", "this is", "based on")):
            name = c; cs = idx + 1; break
    story = []
    if has_photo and name:
        pt = Table([[Paragraph(name.replace("&", "&amp;"), ss["CVNameL"]),
                     Paragraph("[ PHOTO ]<br/>3.5 x 4.5 cm", ss["CVPhoto"])]],
                   colWidths=[12*cm, 4*cm])
        pt.setStyle(TableStyle([
            ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#E8E8E8")),
            ("BOX", (1, 0), (1, 0), 1, colors.HexColor("#999999")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (1, 0), (1, 0), 25),
            ("BOTTOMPADDING", (1, 0), (1, 0), 25),
        ]))
        story.append(pt); story.append(Spacer(1, 10))
    elif name:
        story.append(Paragraph(name.replace("&", "&amp;"), ss["CVName"]))
    for i in range(cs, len(lines)):
        s = lines[i].strip()
        if not s: continue
        if s.lower().startswith(("here is", "below is", "i have", "this is", "based on")):
            continue
        c = re.sub(r"#+\s*", "", s).strip("*").strip()
        if not c: continue
        u = c.upper().replace(":", "")
        safe = c.replace("&", "&amp;").replace("<", "&lt;")
        is_h = s.startswith("#") or (s.startswith("**") and s.endswith("**") and any(k in u for k in SK))
        if not is_h and len(c) >= 3:
            al = [x for x in c if x.isalpha()]
            if al and sum(1 for x in al if x.isupper()) / len(al) > 0.7 and any(k in u for k in SK):
                is_h = True
        if is_h:
            story.append(Paragraph(safe.upper(), ss["CVHead"]))
        elif c.startswith(("-", "*")):
            stripped = safe.lstrip("-* ")
            story.append(Paragraph("• " + stripped, ss["CVBullet"]))
        else:
            story.append(Paragraph(re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe), ss["CVBody"]))
    if ri.get("code") == "IT":
        story.append(Spacer(1, 12))
        story.append(Paragraph(
            "<i>Autorizzo il trattamento dei dati personali ai sensi del GDPR (UE 2016/679).</i>", ss["CVBody"]))
    if ri.get("code") == "IN":
        story.append(Spacer(1, 12))
        story.append(Paragraph(
            "<i>Declaration: I hereby declare that all information provided above is true to the best of my knowledge.</i>", ss["CVBody"]))
    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "<i><font size=7 color='#999999'>AI CV Builder - " +
        datetime.now().strftime("%d %B %Y") + "</font></i>", ss["CVBody"]))
    doc.build(story); buf.seek(0)
    return buf.getvalue()

# ============================================================
# AI PROMPTS
# ============================================================
NO_PRE = "RULES: 1) NO FABRICATION 2) Missing=[Not provided] 3) EXPERIENCE=sum jobs(exclude gaps) 4) Only stated numbers 5) Only stated skills 6) Name first, no preamble 7) ACTUAL years 8) Enhance wording only 9) Strictly follow region format"

def _rb(rn, ri):
    return ("TARGET REGION: " + rn + "\n" +
            "DOCUMENT TYPE: " + ri["term"] + "\n" +
            "LENGTH: " + ri["pages"] + "\n" +
            "PHOTO REQUIREMENT: " + ri["photo"] + "\n" +
            "PERSONAL INFO TO INCLUDE: " + ri["pi"] + "\n" +
            "REQUIRED SECTIONS: " + ri["sec"] + "\n" +
            "STYLE: " + ri["sty"] + "\n" +
            "SPECIAL REQUIREMENTS: " + ri["sp"] + "\n" +
            "ATS: " + ri["ats"] + "\n" +
            "AVOID: " + ri["av"] + "\n\n" +
            "IMPORTANT: Format CV strictly for " + rn + ". If region requires photo, include [PHOTO PLACEHOLDER] at top. Use exact sections listed above. Do NOT use emojis, country flags, or special box-drawing characters in the output.")

def p_gen(jd, info, rn, ri):
    return ("You are an expert CV writer specifically for " + rn + ".\n" + _rb(rn, ri) + "\n" + NO_PRE +
            "\n\nGenerate a " + ri["term"] + " strictly following " + rn + " conventions. Match JD keywords where actual experience aligns." +
            "\n\n---JOB DESCRIPTION---\n" + jd + "\n\n---CANDIDATE INFO---\n" + info)
def p_cmp(cv, jd, rn, ri):
    return ("Expert analyst for " + rn + ".\n" + _rb(rn, ri) +
            "\nProvide: JD MATCH X/100 | REGION COMPLIANCE X/100 | KEY MATCHES | GAPS | SUGGESTIONS" +
            "\n---JD---\n" + jd + "\n---CV---\n" + cv)
def p_imp_cmp(cv, an, jd, rn, ri, kw="", kws="experience"):
    r = ("\nKEYWORD MODE (ATS): MUST place EVERY keyword. Weave naturally." if kws == "force_all"
         else "\nKEYWORD MODE: Only where experience supports. Weave naturally.")
    return ("Expert CV writer for " + rn + ".\n" + _rb(rn, ri) + "\n" + NO_PRE +
            "\nKeep ALL info. Fix issues from analysis. Reformat to strict " + rn + " standard." + r +
            "\n---CV---\n" + cv + "\n---JD---\n" + jd + "\n---ANALYSIS---\n" + an + "\n---KEYWORDS---\n" + kw)
def p_cover(cv, jd, rn, ri):
    return ("Write a cover letter for " + rn + ".\n" + NO_PRE +
            "\nFrom CV+JD. **Bold** key matches. ~350 words. End with KEY MATCHES section. Follow " + rn + " cover letter conventions." +
            "\n---CV---\n" + cv + "\n---JD---\n" + jd + "\n---REGION---\n" + _rb(rn, ri))
def p_ana(cv, rn, ri):
    return ("Expert reviewer for " + rn + ".\n" + _rb(rn, ri) +
            "\nProvide: SCORE X/100 | REGION COMPLIANCE X/100 | STRENGTHS | IMPROVEMENTS | RECOMMENDATIONS" +
            "\n---CV---\n" + cv)
def p_imp(cv, an, rn, ri):
    return ("Expert CV writer for " + rn + ".\n" + _rb(rn, ri) + "\n" + NO_PRE +
            "\nFix ALL issues from analysis. Reformat to strict " + rn + " standards." +
            "\n---CV---\n" + cv + "\n---ANALYSIS---\n" + an)
def p_int(jd):
    return ("Interview coach.\nGenerate 15-20 questions from JD.\nTECHNICAL(10) + BEHAVIOURAL(5) + SITUATIONAL(5)" +
            "\nFor EACH: question | why asked | answer framework (STAR) | key points\nBasic to advanced.\n---JD---\n" + jd)

# ============================================================
# ANTI-REPEAT MOCK PROMPTS — STRONG (Bug fix #1)
# ============================================================
SUBTOPIC_SEEDS = [
    "technical fundamentals & theory",
    "real-world troubleshooting & debugging scenarios",
    "system design & architecture trade-offs",
    "recent industry trends, tools and best practices",
    "team collaboration & cross-functional communication",
    "edge cases, failures, and lessons learned",
    "performance optimisation and scalability",
    "security, compliance and ethics",
    "mentoring, leadership and conflict resolution",
    "career growth, learning curves and pivots",
]

def _build_avoid_block(history_questions):
    """Build STRONG, explicit avoid-list to force AI away from repeats."""
    if not history_questions:
        return ""
    recent = history_questions[-40:]
    listing = "\n".join(f"{i+1}. {q[:180]}" for i, q in enumerate(recent))
    return (
        "\n\n=========================================\n"
        "CRITICAL ANTI-REPEAT INSTRUCTION\n"
        "=========================================\n"
        f"You have ALREADY asked these {len(recent)} questions in previous rounds.\n"
        "YOU MUST GENERATE COMPLETELY DIFFERENT QUESTIONS - DO NOT:\n"
        "  - Repeat any of them word-for-word\n"
        "  - Paraphrase them\n"
        "  - Ask the same concept with different wording\n"
        "  - Use similar examples or scenarios\n\n"
        "ALREADY ASKED (FORBIDDEN - generate completely fresh ones):\n"
        + listing +
        "\n\nIf you cannot think of new questions, dig DEEPER into sub-topics, "
        "advanced scenarios, edge cases, modern tools, recent industry trends, "
        "specific real-world situations. Be creative and diverse.\n"
        "=========================================\n"
    )

def p_mock_jd(jd, n=15, exp_level="Mid-level (2-5 yrs)", history_questions=None):
    nonce = uuid.uuid4().hex[:12]
    seed = random.randint(10000, 99999)
    focus = random.sample(SUBTOPIC_SEEDS, k=min(4, len(SUBTOPIC_SEEDS)))
    avoid_block = _build_avoid_block(history_questions or [])
    return ("You are a CREATIVE and VARIED interview coach.\n"
            "EXPERIENCE LEVEL OF CANDIDATE: " + exp_level + "\n"
            "Tailor difficulty + scenarios for this level. For Students/Freshers focus on fundamentals, "
            "academic projects, internships. For Senior/Lead focus on architecture, leadership, trade-offs.\n\n"
            "ROTATE FOCUS - this round must emphasise these specific sub-topics:\n"
            "- " + "\n- ".join(focus) + "\n\n"
            "Generate " + str(n) + " FRESH, UNIQUE, NON-OVERLAPPING questions from JD.\n"
            "60% Technical + 25% Behavioural + 15% Situational.\n"
            "Vary phrasing, scenario, depth. Use modern industry context.\n"
            "Session randomness token: " + nonce + "-" + str(seed) + "\n\n"
            "OUTPUT FORMAT (strict):\nQ1: [question]\nQ2: [question]\n...Only the questions."
            + avoid_block + "\n---JD---\n" + jd)

def p_mock_dom(dom, n=15, exp_level="Mid-level (2-5 yrs)", history_questions=None):
    nonce = uuid.uuid4().hex[:12]
    seed = random.randint(10000, 99999)
    focus = random.sample(SUBTOPIC_SEEDS, k=min(4, len(SUBTOPIC_SEEDS)))
    avoid_block = _build_avoid_block(history_questions or [])
    return ("You are a CREATIVE and VARIED interview coach for **" + dom + "**.\n"
            "EXPERIENCE LEVEL OF CANDIDATE: " + exp_level + "\n"
            "Tailor difficulty for this level. For Students/Freshers focus on fundamentals, theory, college projects. "
            "For Senior/Lead focus on architecture, leadership, trade-offs, mentoring.\n\n"
            "ROTATE FOCUS - this round must emphasise these specific sub-topics of " + dom + ":\n"
            "- " + "\n- ".join(focus) + "\n\n"
            "Generate " + str(n) + " FRESH, UNIQUE, NON-OVERLAPPING questions SPECIFICALLY for " + dom + ".\n"
            "Stay tightly focused on " + dom + ".\n"
            "60% Technical + 25% Behavioural + 15% Situational.\n"
            "Vary phrasing, scenario, depth. Use modern industry context.\n"
            "Session randomness token: " + nonce + "-" + str(seed) + "\n\n"
            "OUTPUT FORMAT (strict):\nQ1: [question]\nQ2: [question]\n...Only the questions."
            + avoid_block)

def p_mock_eval(q, a, exp_level="Mid-level (2-5 yrs)"):
    return ("Interview coach. Evaluate this answer for a " + exp_level + " candidate.\n"
            "Adjust expectations to that level.\n"
            "Question: " + q + "\nAnswer: " + a +
            "\n\nProvide:\nSCORE X/10 | STRENGTHS | WEAKNESSES | IDEAL ANSWER (STAR) | 3 TIPS | CONFIDENCE: Low/Med/High")

def p_coach(t, c):
    return ("Career coach: " + t + "\nContext: " + c + "\nInsights, Action plan, Pitfalls, Resources, Timeline")
def p_multi(cv, jds, rn, ri):
    return ("Analyst for " + rn + ".\n" + _rb(rn, ri) +
            "\nCompare CV vs MULTIPLE JDs.\nFor EACH: Title | Match X/100 | Matches | Gaps | Keywords\nOVERALL: best fit + why." +
            "\n---CV---\n" + cv + "\n\n" + jds)

INFO_T = """Only provide REAL information.
Full Name:
Email:
Phone:
Location:
LinkedIn:
--- WORK EXPERIENCE ---
### Job 1:
Title:
Company:
From:
To:
What you did:
### Job 2:
Title:
Company:
From:
To:
What you did:
--- EDUCATION ---
Qualification:
Institution:
Year:
--- SKILLS ---
Technical:
Soft Skills:
Languages:
--- CERTIFICATIONS ---
Name:
Year:
"""

def _parse_questions(text):
    """SUPER ROBUST parser — catches almost any AI question format. Rejects preamble."""
    if not text: return []
    out = []
    cleaned = re.sub(r'\*+', '', text)
    cleaned = re.sub(r'^#+\s*', '', cleaned, flags=re.MULTILINE)

    # Lines/phrases that are AI preamble, NOT real questions
    INTRO_BLOCKERS = [
        "here are", "here is", "below are", "below is",
        "these are", "this is", "i have", "i'll provide",
        "let me", "as an", "as a", "based on", "okay",
        "sure", "of course", "happy to", "absolutely",
        "technical questions", "behavioural questions",
        "behavioral questions", "situational questions",
        "section", "category", "round", "part",
    ]
    SECTION_HEADERS = {
        "technical", "behavioural", "behavioral", "situational",
        "technical questions", "behavioural questions",
        "behavioral questions", "situational questions",
    }

    patterns = [
        r"^[\W_]*Q\s*\.?\s*\d+\s*[:.\)\-\]]\s*(.+)$",
        r"^[\W_]*\d+\s*[:.\)\-\]]\s*(.+)$",
        r"^[\W_]*Question\s*\d+\s*[:.\)\-\]]\s*(.+)$",
    ]

    for line in cleaned.split("\n"):
        line = line.strip()
        if not line: continue
        if line.startswith(("---", "===", "***", "###", "___")): continue
        low = line.lower().strip(" :*-#")
        if low in SECTION_HEADERS: continue
        if any(low.startswith(b) for b in INTRO_BLOCKERS): continue

        matched = False
        for pat in patterns:
            m = re.match(pat, line, flags=re.I)
            if m:
                q = m.group(1).strip(" :.-_*[]")
                # Re-check for preamble after pattern strip
                q_low = q.lower()
                if any(q_low.startswith(b) for b in INTRO_BLOCKERS): continue
                if len(q) > 10:
                    out.append(q)
                matched = True
                break

        # Fallback for line ending with ? AND has reasonable length
        if not matched and line.endswith("?") and len(line) > 20:
            if not line.isupper() and not line.startswith(("#", "-", "*", "=")):
                cleaned_q = re.sub(r"^[\W_\d]+", "", line).strip()
                if len(cleaned_q) > 15:
                    out.append(cleaned_q)
                elif len(line) > 15:
                    out.append(line)

    return out


# ============================================================
# STREAMLIT CONFIG + CSS
# ============================================================
st.set_page_config(page_title="AI CV Builder", page_icon="📄",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
    .home-card {
        padding: 24px; border-radius: 10px; margin-bottom: 12px;
        min-height: 110px; cursor: pointer; transition: transform 0.15s, box-shadow 0.15s;
        border: 2px solid transparent;
    }
    .home-card:hover { transform: translateY(-2px); box-shadow: 0 6px 16px rgba(0,0,0,0.15); }
    .home-card h3 { margin-top: 0; color: #1a1a2e; }
    .home-card p  { color: #444; margin: 0; }
    .home-card .tap-hint { font-size: 12px; color: #1976d2; margin-top: 8px; font-weight: 600; }
    .bg-green  { background:#d4edda; border-color:#c3e6cb; }
    .bg-yellow { background:#fff3cd; border-color:#ffeaa7; }
    .bg-blue   { background:#d1ecf1; border-color:#bee5eb; }
    .bg-purple { background:#e8daef; border-color:#d7bce8; }
    .hero { background:linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); color:white; padding:30px; border-radius:8px; margin-bottom:20px; }
    .hero h1 { color:white; margin:0 0 10px 0; }
    .hero p { color:#b0c4de; margin:0; }
    .stats-bar { background:#e8eaf6; padding:12px; border-radius:8px; margin-top:12px; text-align:center; }
    .region-warning { background:#fff3cd; padding:10px; border-radius:6px; border-left:4px solid #ff9800; margin-bottom:10px; font-size:14px; }
    div[data-testid="stDownloadButton"] button { font-weight: 600; }
    .onboard {
        background: linear-gradient(135deg, #4CAF50 0%, #2E7D32 100%);
        color: white; padding: 24px; border-radius: 12px;
        margin-bottom: 20px; box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    .onboard h2 { color: white; margin: 0 0 12px 0; font-size: 22px; }
    .onboard p  { color: #E8F5E9; font-size: 16px; line-height: 1.9; margin: 0; }
    .onboard a  { color: #FFEB3B; font-weight: bold; text-decoration: underline; }
    .nokey-banner {
        background: linear-gradient(135deg, #FF9800 0%, #F57C00 100%);
        color: white; padding: 18px; border-radius: 10px; margin-bottom: 16px;
    }
    .nokey-banner h3 { color: white; margin: 0 0 8px 0; }
    .nokey-banner p  { color: #FFF3E0; margin: 0; line-height: 1.7; }
    section[data-testid="stSidebar"] button[kind="header"] {
        position: fixed !important; top: 10px !important; z-index: 999999 !important;
    }
    [data-testid="collapsedControl"] {
        position: fixed !important; top: 10px !important; left: 10px !important;
        z-index: 999999 !important; background: rgba(255,255,255,0.95) !important;
        border-radius: 6px !important; padding: 4px !important;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# PWA: Make installable on iPhone/Android home screen
# ============================================================
st.markdown("""
<link rel="manifest" href="./app/static/manifest.json">
<meta name="theme-color" content="#4CAF50">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="CV Builder">
<meta name="application-name" content="CV Builder">
<link rel="apple-touch-icon" href="./app/static/icon-192.png">
<link rel="icon" type="image/png" sizes="192x192" href="./app/static/icon-192.png">
<link rel="icon" type="image/png" sizes="512x512" href="./app/static/icon-512.png">
""", unsafe_allow_html=True)

# ============================================================
# MOBILE: auto-close sidebar after tab pick
# ============================================================
components.html("""
<script>
(function() {
    const W = window.parent;
    const D = W.document;
    function isMobile() { return W.innerWidth < 992; }
    function findCloseBtn() {
        return D.querySelector('section[data-testid="stSidebar"] button[kind="header"]')
            || D.querySelector('[data-testid="stSidebarCollapseButton"]')
            || D.querySelector('[data-testid="collapsedControl"]');
    }
    function isSidebarOpen() {
        const sb = D.querySelector('section[data-testid="stSidebar"]');
        if (!sb) return false;
        const rect = sb.getBoundingClientRect();
        return rect.width > 50;
    }
    function tryClose() {
        if (!isMobile()) return;
        let tries = 0;
        const iv = setInterval(() => {
            if (!isSidebarOpen() || tries > 12) { clearInterval(iv); return; }
            const btn = findCloseBtn();
            if (btn) btn.click();
            tries++;
        }, 80);
    }
    function attach() {
        const radios = D.querySelectorAll('section[data-testid="stSidebar"] input[type="radio"]');
        radios.forEach(r => {
            if (r._autoCloseAttached) return;
            r._autoCloseAttached = true;
            r.addEventListener('change', () => setTimeout(tryClose, 150));
            r.addEventListener('click', () => setTimeout(tryClose, 200));
        });
        const labels = D.querySelectorAll('section[data-testid="stSidebar"] label');
        labels.forEach(l => {
            if (l._autoCloseAttached) return;
            l._autoCloseAttached = true;
            l.addEventListener('click', () => setTimeout(tryClose, 200));
        });
    }
    attach();
    setInterval(attach, 1000);
    const obs = new MutationObserver(attach);
    obs.observe(D.body, {childList: true, subtree: true});
})();
</script>
""", height=0)

# ============================================================
# NAVIGATION STATE
# ============================================================
PAGES = ["🏠 Home", "📝 Generate CV", "🔍 CV vs JD", "📊 CV Analysis",
         "📑 Multi-JD Compare", "🎤 Interview Prep", "🎙️ Mock Interview",
         "🧑‍💼 Coaching", "📚 My Library", "⚙️ Settings"]

PAGE_SLUGS = {
    "home": "🏠 Home", "generate": "📝 Generate CV", "compare": "🔍 CV vs JD",
    "analysis": "📊 CV Analysis", "multi": "📑 Multi-JD Compare",
    "prep": "🎤 Interview Prep", "mock": "🎙️ Mock Interview",
    "coach": "🧑‍💼 Coaching", "library": "📚 My Library", "settings": "⚙️ Settings",
}

if "page" not in st.session_state:
    st.session_state.page = "🏠 Home"

qp = st.query_params
if "nav" in qp:
    target_slug = qp.get("nav")
    if target_slug in PAGE_SLUGS:
        st.session_state.page = PAGE_SLUGS[target_slug]
    st.query_params.clear()
    st.rerun()

# ============================================================
# SIDEBAR — Gemini + OpenAI + Claude
# ============================================================
with st.sidebar:
    st.markdown("# 📄 AI CV Builder")
    st.markdown("---")
    idx = PAGES.index(st.session_state.page) if st.session_state.page in PAGES else 0
    selected = st.radio("**Navigation**", PAGES, index=idx, label_visibility="collapsed",
                        key="nav_radio")
    if selected != st.session_state.page:
        st.session_state.page = selected
        st.rerun()
    page = st.session_state.page

    st.markdown("---")
    st.markdown("**🤖 AI Provider**")
    provider = st.radio("Provider",
                        ["Gemini (FREE)", "OpenAI (Paid)", "Claude (Paid)"],
                        label_visibility="collapsed", key="prov_radio")
    if "OpenAI" in provider:
        prov = "openai"; models = OPENAI_MODELS
    elif "Claude" in provider:
        prov = "claude"; models = CLAUDE_MODELS
    else:
        prov = "gemini"; models = GEMINI_MODELS

    model_display = st.selectbox("Model", [m[0] + " (" + m[1] + ")" for m in models],
                                 key="model_sel")
    model_id = model_display.split(" (")[0]
    api_key = st.text_input("🔑 API Key", type="password",
                            placeholder="Paste your API key", key="api_key_in")

    if prov == "gemini":
        st.markdown("🔗 [Get FREE Gemini key](https://aistudio.google.com/apikey)")
    elif prov == "openai":
        st.markdown("🔗 https://platform.openai.com/api-keys")
    else:
        st.markdown("🔗 https://console.anthropic.com/settings/keys")

    st.markdown("---")
    st.markdown("**🌍 Target Region** ⭐")
    region = st.selectbox("Region", RL, label_visibility="collapsed", key="region_select",
                          help="CV format adapts to region")
    region_name = region.split(" ", 1)[1] if " " in region else region
    st.caption("📍 **" + region_name + "** format will be used")
    st.caption("SSL: " + SSL_M)

# ============================================================
# HELPERS
# ============================================================
def call_ai(prompt, temperature=0.7):
    if not api_key:
        st.markdown("""
        <div class="nokey-banner">
            <h3>🔑 API key missing</h3>
            <p>You haven't entered your API key in the left sidebar yet.<br>
            <b>👉 What to do:</b><br>
            1. Open the sidebar (tap the <b>»</b> arrow top-left on phone).<br>
            2. Choose a provider — <b>Gemini (FREE)</b> is recommended.<br>
            3. Paste your API key in the 🔑 box.<br>
            4. Don't have one? Get a free Gemini key at
            <a href="https://aistudio.google.com/apikey" target="_blank" style="color:#FFEB3B">
            aistudio.google.com/apikey</a> (30 seconds).</p>
        </div>
        """, unsafe_allow_html=True)
        return None
    with st.spinner("🤖 AI is thinking..."):
        return ai_call(prompt, api_key, prov, model_id, temperature=temperature)

def download_buttons(text, region, suffix="Document", name_override=None):
    if not text: return
    name = name_override or extract_name(text)
    base_name = f"{name}_{suffix}"
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("💾 Save (.txt)", strip_emojis(text), base_name + ".txt",
                           mime="text/plain", use_container_width=True,
                           key="dl_txt_" + base_name)
    with col2:
        if HAS_DOCX:
            docx_data = make_docx(text, region)
            if docx_data:
                st.download_button("📄 Word (.docx)", docx_data, base_name + ".docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="dl_docx_" + base_name)
        else:
            st.button("📄 Word (need python-docx)", disabled=True,
                      use_container_width=True, key="dl_docx_dis_" + base_name)
    with col3:
        if HAS_RL:
            pdf_data = make_pdf(text, region)
            if pdf_data:
                st.download_button("📑 PDF (.pdf)", pdf_data, base_name + ".pdf",
                    mime="application/pdf", use_container_width=True,
                    key="dl_pdf_" + base_name)
        else:
            st.button("📑 PDF (need reportlab)", disabled=True,
                      use_container_width=True, key="dl_pdf_dis_" + base_name)

def region_note(stored_region):
    if stored_region and stored_region != region:
        st.markdown('<div class="region-warning">⚠️ Region changed from <b>' + stored_region +
                    '</b> to <b>' + region + '</b>. Click Generate/Improve again.</div>',
                    unsafe_allow_html=True)


# ============================================================
# PAGE: HOME
# ============================================================
if page == "🏠 Home":
    st.markdown(
        '<div class="hero"><h1>📄 AI CV Builder</h1>'
        '<p>Gemini (FREE) + OpenAI + Claude • Region-aware • PDF/Word Export • 130+ Mock Interview Domains • Voice Input + AI Polish<br>'
        'Anti-hallucination: AI only uses YOUR data.</p></div>',
        unsafe_allow_html=True)

    if not api_key:
        st.markdown("""
        <div class="onboard">
            <h2>🆓 First time? Get a FREE Gemini API key in 30 seconds:</h2>
            <p>
                <b>Step 1:</b> Click → istudio.google.com/apikey" target="_blank">https://aistudio.google.com/apikey</a><br>
                <b>Step 2:</b> Sign in with your Google account<br>
                <b>Step 3:</b> Click 'Create API Key' → copy it → paste in sidebar<br>
                <br>
                <i>It's 100% free with generous daily quota.</i>
            </p>
        </div>
        """, unsafe_allow_html=True)

    cards = [
        ("📝 Generate CV",     "Create a region-formatted CV from any JD.",         "bg-green",  "generate"),
        ("🔍 CV vs JD",        "Compare CV vs JD. Keywords, gaps, cover letter.",   "bg-yellow", "compare"),
        ("📊 CV Analysis",     "AI scores ATS compliance. Improve & export.",       "bg-blue",   "analysis"),
        ("📑 Multi-JD Compare","Compare CV against multiple JDs at once.",          "bg-purple", "multi"),
        ("🎤 Interview Prep",  "15-20 questions + STAR frameworks from any JD.",    "bg-green",  "prep"),
        ("🎙️ Mock Interview", "Voice + AI polish + 130+ domains.",                  "bg-yellow", "mock"),
        ("🧑‍💼 Coaching",     "Career advice on 10+ topics.",                       "bg-blue",   "coach"),
        ("📚 My Library",      "Save question sets & feedback. Export as JSON.",    "bg-green",  "library"),
        ("⚙️ Settings",        "Providers, regions, mobile tips.",                  "bg-purple", "settings"),
    ]
    cols = st.columns(2, gap="medium")
    for i, (title, desc, css_class, slug) in enumerate(cards):
        with cols[i % 2]:
            st.markdown(
                f'?nav={slug}" style="text-decoration:none;color:inherit;">'
                f'<div class="home-card {css_class}">'
                f'<h3>{title}</h3><p>{desc}</p>'
                f'<div class="tap-hint">👉 Tap to open</div>'
                f'</div></a>',
                unsafe_allow_html=True)

    st.markdown(
        '<div class="stats-bar">🌍 <b>14 Regions</b> &nbsp;•&nbsp; 🤖 <b>Gemini + OpenAI + Claude</b> '
        '&nbsp;•&nbsp; 🎤 <b>Voice + AI Polish</b> &nbsp;•&nbsp; 📑 <b>PDF / Word</b> '
        '&nbsp;•&nbsp; 🎙️ <b>130+ Mock Domains</b></div>', unsafe_allow_html=True)
    st.markdown("---")
    st.info("💡 **Currently set to:** " + region + " — change in the sidebar anytime.")

# ============================================================
# PAGE: GENERATE CV
# ============================================================
elif page == "📝 Generate CV":
    st.title("📝 Generate CV")
    st.caption("Currently generating for: **" + region + "** - " + RD_[region]["photo"])
    region_note(st.session_state.get("gen_region"))
    jd_gen = st.text_area("📋 Job Description", height=200, key="gen_jd")
    info_gen = st.text_area("👤 Your Information", height=400, key="gen_info",
                            value=st.session_state.get("gen_info_val", INFO_T))
    if st.button("✨ Generate CV", type="primary", use_container_width=True):
        if not jd_gen.strip(): st.warning("Please paste a Job Description.")
        elif not info_gen.strip(): st.warning("Please fill in your information.")
        else:
            result = call_ai(p_gen(jd_gen, info_gen, region, RD_[region]))
            if result:
                st.session_state["gen_result"] = result
                st.session_state["gen_region"] = region
                st.session_state["gen_info_val"] = info_gen
    if "gen_result" in st.session_state:
        st.markdown("### ✨ Generated CV (formatted for " +
                    st.session_state.get("gen_region", region) + ")")
        st.text_area("Result", st.session_state["gen_result"], height=500,
                     key="gen_out", label_visibility="collapsed")
        st.markdown("**Download:**")
        download_buttons(st.session_state["gen_result"],
                         st.session_state.get("gen_region", region), "CV")

# ============================================================
# PAGE: CV vs JD
# ============================================================
elif page == "🔍 CV vs JD":
    st.title("🔍 CV vs Job Description")
    st.caption("Currently using region: **" + region + "**")
    region_note(st.session_state.get("cmp_region"))
    col1, col2 = st.columns(2, gap="medium")
    with col1:
        st.markdown("### 📄 Your CV")
        cv_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"], key="cmp_cv_file")
        cv_text_input = st.text_area("Or paste CV", height=300, key="cmp_cv_input")
        cv_text = read_file(cv_file) if cv_file else cv_text_input
    with col2:
        st.markdown("### 📋 Job Description")
        jd_cmp = st.text_area("Paste JD", height=300, key="cmp_jd", label_visibility="collapsed")
    if st.button("🔍 Compare CV vs JD", type="primary", use_container_width=True):
        if not cv_text.strip() or not jd_cmp.strip():
            st.warning("Please provide both CV and JD.")
        else:
            result = call_ai(p_cmp(cv_text, jd_cmp, region, RD_[region]))
            if result:
                st.session_state["cmp_result"] = result
                st.session_state["cmp_cv"]     = cv_text
                st.session_state["cmp_jd"]     = jd_cmp
                st.session_state["cmp_region"] = region
    if "cmp_result" in st.session_state:
        st.markdown("### 📊 Comparison Result")
        st.text_area("Result", st.session_state["cmp_result"], height=400,
                     key="cmp_out", label_visibility="collapsed")
        cv_name = extract_name(st.session_state.get("cmp_cv", ""))
        download_buttons(st.session_state["cmp_result"], region,
                         "JD_Comparison", name_override=cv_name)

        st.markdown("---")
        st.markdown("### ✨ Improve CV with Keywords")
        col_k1, col_k2 = st.columns(2)
        with col_k1:
            kw_strength = st.radio("Keyword Strength", ["Experience-based", "Force ALL"],
                                   key="kw_str", horizontal=True)
        with col_k2:
            kw_mode = st.radio("Keywords from", ["Auto from JD", "Manual"],
                               key="kw_mode", horizontal=True)
        manual_kw = ""
        if kw_mode == "Manual":
            manual_kw = st.text_input("Enter keywords (comma-separated)",
                                      placeholder="Docker, Kubernetes, CI/CD")
        if st.button("✨ Improve CV", type="primary", use_container_width=True):
            strength = "force_all" if kw_strength == "Force ALL" else "experience"
            result = call_ai(p_imp_cmp(st.session_state["cmp_cv"], st.session_state["cmp_result"],
                                       st.session_state["cmp_jd"], region, RD_[region],
                                       manual_kw, strength))
            if result:
                st.session_state["cmp_improved"] = result
                st.session_state["cmp_region"]   = region
        if "cmp_improved" in st.session_state:
            st.markdown("#### ✨ Improved CV")
            st.text_area("Improved", st.session_state["cmp_improved"], height=500,
                         key="cmp_imp_out", label_visibility="collapsed")
            download_buttons(st.session_state["cmp_improved"],
                             st.session_state.get("cmp_region", region),
                             "Improved_CV")

        st.markdown("---")
        st.markdown("### 📨 Cover Letter")
        if st.button("📨 Generate Cover Letter", type="primary", use_container_width=True):
            cv_src = st.session_state.get("cmp_improved", st.session_state.get("cmp_cv", ""))
            result = call_ai(p_cover(cv_src, st.session_state["cmp_jd"], region, RD_[region]))
            if result: st.session_state["cmp_cl"] = result
        if "cmp_cl" in st.session_state:
            st.text_area("Cover Letter", st.session_state["cmp_cl"], height=400,
                         key="cmp_cl_out", label_visibility="collapsed")
            cv_name = extract_name(st.session_state.get("cmp_improved",
                                   st.session_state.get("cmp_cv", st.session_state["cmp_cl"])))
            download_buttons(st.session_state["cmp_cl"], region,
                             "Cover_Letter", name_override=cv_name)

# ============================================================
# PAGE: CV ANALYSIS
# ============================================================
elif page == "📊 CV Analysis":
    st.title("📊 CV Analysis")
    st.caption("Analysing against **" + region + "** standards.")
    region_note(st.session_state.get("ana_region"))
    cv_ana_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"], key="ana_file")
    cv_ana_input = st.text_area("Or paste CV", height=300, key="ana_cv_input")
    cv_ana = read_file(cv_ana_file) if cv_ana_file else cv_ana_input
    st.info("**Score guide:** 90-100 Excellent | 75-89 Good | 60-74 Average | <60 Needs work")
    if st.button("📊 Analyse CV", type="primary", use_container_width=True):
        if not cv_ana.strip(): st.warning("Please provide a CV.")
        else:
            result = call_ai(p_ana(cv_ana, region, RD_[region]))
            if result:
                st.session_state["ana_result"] = result
                st.session_state["ana_cv"]     = cv_ana
                st.session_state["ana_region"] = region
    if "ana_result" in st.session_state:
        st.markdown("### 📊 Analysis")
        st.text_area("Analysis", st.session_state["ana_result"], height=400,
                     key="ana_out", label_visibility="collapsed")
        cv_name = extract_name(st.session_state.get("ana_cv", ""))
        download_buttons(st.session_state["ana_result"], region,
                         "CV_Analysis", name_override=cv_name)
        st.markdown("---")
        if st.button("✨ Improve CV", type="primary", use_container_width=True):
            result = call_ai(p_imp(st.session_state["ana_cv"], st.session_state["ana_result"],
                                   region, RD_[region]))
            if result:
                st.session_state["ana_improved"] = result
                st.session_state["ana_region"]   = region
        if "ana_improved" in st.session_state:
            st.markdown("### ✨ Improved CV")
            st.text_area("Improved", st.session_state["ana_improved"], height=500,
                         key="ana_imp_out", label_visibility="collapsed")
            download_buttons(st.session_state["ana_improved"],
                             st.session_state.get("ana_region", region),
                             "Improved_CV")

# ============================================================
# PAGE: MULTI-JD COMPARE
# ============================================================
elif page == "📑 Multi-JD Compare":
    st.title("📑 Multi-JD Compare")
    st.caption("Using region: **" + region + "**")
    cv_multi_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"], key="multi_file")
    cv_multi_input = st.text_area("Or paste CV", height=200, key="multi_cv_input")
    cv_multi = read_file(cv_multi_file) if cv_multi_file else cv_multi_input
    n_jds = st.number_input("Number of JDs to compare", min_value=2, max_value=10,
                            value=2, key="multi_n")
    jd_texts = []
    for i in range(int(n_jds)):
        jd_texts.append(st.text_area("📋 JD " + str(i + 1), height=150,
                                     key="multi_jd_" + str(i)))
    if st.button("📑 Compare All JDs", type="primary", use_container_width=True):
        if not cv_multi.strip():
            st.warning("Please provide a CV.")
        else:
            combined = "\n\n".join(
                ["--- JD " + str(i + 1) + " ---\n" + t
                 for i, t in enumerate(jd_texts) if t.strip()])
            if not combined.strip():
                st.warning("Please fill in at least one JD.")
            else:
                result = call_ai(p_multi(cv_multi, combined, region, RD_[region]))
                if result:
                    st.session_state["multi_result"] = result
                    st.session_state["multi_cv"]     = cv_multi
    if "multi_result" in st.session_state:
        st.markdown("### 📊 Comparison Results")
        st.text_area("Results", st.session_state["multi_result"], height=500,
                     key="multi_out", label_visibility="collapsed")
        cv_name = extract_name(st.session_state.get("multi_cv", ""))
        download_buttons(st.session_state["multi_result"], region,
                         "Multi_JD_Results", name_override=cv_name)

# ============================================================
# PAGE: INTERVIEW PREP
# ============================================================
elif page == "🎤 Interview Prep":
    st.title("🎤 Interview Preparation")
    st.info("💡 Questions come from AI knowledge of interview patterns.")
    jd_intv = st.text_area("📋 Job Description", height=250, key="intv_jd")
    if st.button("🎯 Generate Questions", type="primary", use_container_width=True):
        if not jd_intv.strip():
            st.warning("Please paste a JD.")
        else:
            result = call_ai(p_int(jd_intv))
            if result:
                st.session_state["intv_result"] = result
    if "intv_result" in st.session_state:
        st.markdown("### 🎯 Questions + Answer Frameworks")
        st.text_area("Questions", st.session_state["intv_result"], height=600,
                     key="intv_out", label_visibility="collapsed")
        download_buttons(st.session_state["intv_result"], region,
                         "Interview_Questions", name_override="Candidate")

# ============================================================
# PAGE: MOCK INTERVIEW v5 — fixed Polish/Use Raw/Prev/Next + bigger transcript box
# ============================================================
elif page == "🎙️ Mock Interview":
    st.title("🎙️ Mock Interview")
    st.caption("Pick Q → speak → polish → evaluate. Progress saved per question!")

    if "mock_history" not in st.session_state:
        st.session_state["mock_history"] = {}
    if "q_state" not in st.session_state:
        st.session_state["q_state"] = {}

    def _save_current_q():
        idx = st.session_state.get("mock_q_idx")
        if idx is None or "mock_q_list" not in st.session_state:
            return
        st.session_state["q_state"][idx] = {
            "transcript": st.session_state.get("raw_voice_in", ""),
            "polished": st.session_state.get("polished_preview", ""),
            "answer": st.session_state.get("mock_a_in", ""),
            "feedback": st.session_state.get("mock_feedback", ""),
        }

    def _load_q(idx):
        state = st.session_state["q_state"].get(idx, {})
        st.session_state["raw_voice_in"] = state.get("transcript", "")
        st.session_state["mock_a_in"] = state.get("answer", "")
        if state.get("polished"):
            st.session_state["polished_preview"] = state["polished"]
        else:
            st.session_state.pop("polished_preview", None)
        if state.get("feedback"):
            st.session_state["mock_feedback"] = state["feedback"]
        else:
            st.session_state.pop("mock_feedback", None)

    def _switch_to(new_idx):
        """Switch to new Q — save current, load new, set ONLY mock_q_idx (not picker)."""
        _save_current_q()
        st.session_state["mock_q_idx"] = new_idx
        _load_q(new_idx)

    # ---- Source picker ----
    with st.container(border=True):
        st.markdown("### 📋 Question Source")
        source = st.radio("Source", ["📋 From Job Description", "🎯 By Domain / Field"],
                          horizontal=True, key="mock_src", label_visibility="collapsed")

        mock_jd = ""
        category = ""
        mock_domain = ""

        if source == "📋 From Job Description":
            mock_jd = st.text_area("Paste JD", height=200, key="mock_jd")
        elif source == "🎯 By Domain / Field":
            cat_col, dom_col = st.columns([1, 2])
            with cat_col:
                category = st.selectbox("Category", list(DOMAIN_GROUPS.keys()), key="mock_cat")
            with dom_col:
                mock_domain = st.selectbox("Domain / Field", DOMAIN_GROUPS[category], key="mock_dom")
            st.caption(f"📚 **{len(DOMAIN_GROUPS[category])}** option(s) under **{category}** "
                       f"• Total: **{sum(len(v) for v in DOMAIN_GROUPS.values())}**")

        exp_level = st.selectbox("🎯 Experience Level", EXP_LEVELS, index=2, key="mock_exp")
        n_q = st.radio("Number of questions", ["10", "15", "20"], index=1,
                       horizontal=True, key="mock_n")

    if source == "📋 From Job Description":
        history_key = f"JD::{exp_level}::{(mock_jd[:80] or 'na')}"
    else:
        history_key = f"DOM::{category}::{mock_domain}::{exp_level}"

    hist_count = len(st.session_state["mock_history"].get(history_key, []))
    if hist_count > 0:
        st.caption(f"📊 Anti-repeat history: **{hist_count} questions tracked**")

    sig = f"{source}|{category}|{mock_domain}|{exp_level}|{n_q}|{mock_jd[:100]}"
    if st.session_state.get("mock_sig") and st.session_state["mock_sig"] != sig and \
       "mock_questions" in st.session_state:
        st.warning("⚠️ Settings changed — old questions hidden. Tap Generate.")
        for k in ["mock_questions", "mock_q_list", "mock_q_idx", "q_state",
                  "raw_voice_in", "mock_a_in", "mock_feedback", "polished_preview"]:
            st.session_state.pop(k, None)

    col_g1, col_g2, col_g3 = st.columns([2, 2, 1])
    with col_g1:
        gen_clicked = st.button("🎯 Generate Questions", type="primary",
                                use_container_width=True, key="mock_gen_btn")
    with col_g2:
        regen_clicked = st.button("🔄 Regenerate (different)", use_container_width=True,
                                  key="mock_regen_btn")
    with col_g3:
        clear_clicked = st.button("🧹 Clear", use_container_width=True,
                                  key="mock_clear_btn")

    if clear_clicked:
        st.session_state["mock_history"].pop(history_key, None)
        for k in ["mock_questions", "mock_q_list", "mock_q_idx", "q_state",
                  "raw_voice_in", "mock_a_in", "mock_feedback", "polished_preview"]:
            st.session_state.pop(k, None)
        st.success("✅ History cleared.")
        st.rerun()

    if gen_clicked or regen_clicked:
        history = st.session_state["mock_history"].get(history_key, [])
        if source == "📋 From Job Description":
            if not mock_jd.strip():
                st.warning("Please paste a JD first.")
                st.stop()
            prompt = p_mock_jd(mock_jd, int(n_q), exp_level, history_questions=history)
            meta_domain = "From JD"
        else:
            if not mock_domain:
                st.warning("Please pick a category and domain.")
                st.stop()
            prompt = p_mock_dom(mock_domain, int(n_q), exp_level, history_questions=history)
            meta_domain = f"{category} → {mock_domain}"

        result = call_ai(prompt, temperature=1.0)
        if result:
            st.session_state["mock_questions"] = result
            st.session_state["mock_sig"] = sig
            st.session_state["mock_meta"] = {
                "domain": meta_domain, "level": exp_level, "n": n_q,
                "ts": datetime.now().strftime("%Y-%m-%d %H:%M"),
            }
            new_qs = _parse_questions(result)
            if new_qs:
                st.session_state["mock_history"].setdefault(history_key, []).extend(new_qs)
                st.session_state["mock_q_list"] = new_qs
                st.session_state["mock_q_idx"] = 0
                st.session_state["q_state"] = {}
                for k in ["raw_voice_in", "mock_a_in", "mock_feedback", "polished_preview"]:
                    st.session_state.pop(k, None)
                st.success(f"✅ Generated {len(new_qs)} questions (history: "
                           f"{len(st.session_state['mock_history'][history_key])})")
            else:
                # Stricter fallback: ONLY lines that look like real questions
                fallback_qs = []
                INTRO_BLOCKERS = [
                    "here are", "here is", "below are", "below is",
                    "these are", "this is", "i have", "i'll provide",
                    "let me", "as an", "as a", "based on", "okay",
                    "sure", "of course", "happy to", "absolutely",
                    "technical questions", "behavioural questions",
                    "behavioral questions", "situational questions",
                    "section", "category", "round", "part",
                ]
                for ln in result.split("\n"):
                    ln = ln.strip().strip("*").strip("-").strip("#").strip()
                    if len(ln) < 20 or len(ln) > 500: continue
                    if ln.isupper(): continue
                    low = ln.lower()
                    # Reject intro/preamble lines
                    if any(low.startswith(b) for b in INTRO_BLOCKERS): continue
                    # Must look like a question OR have numbering
                    has_question_mark = "?" in ln
                    starts_with_num = bool(re.match(r"^[\d\W]*\d+", ln))
                    if not (has_question_mark or starts_with_num): continue
                    # Strip leading numbering noise
                    cleaned_q = re.sub(r"^[\W_\d]*\d+\s*[:.\)\-\]]\s*", "", ln).strip()
                    cleaned_q = cleaned_q.strip("*").strip()
                    if len(cleaned_q) > 15:
                        fallback_qs.append(cleaned_q)

                if len(fallback_qs) >= 3:
                    fallback_qs = fallback_qs[:int(n_q)]
                    st.session_state["mock_history"].setdefault(history_key, []).extend(fallback_qs)
                    st.session_state["mock_q_list"] = fallback_qs
                    st.session_state["mock_q_idx"] = 0
                    st.session_state["q_state"] = {}
                    st.info(f"ℹ️ Used fallback parser — got {len(fallback_qs)} questions. "
                            "Format was unusual but questions are usable.")
                else:
                    st.warning("⚠️ AI returned unusual format. Tap 🧹 Clear and retry. "
                               "If it keeps happening, switch to **Gemini 2.5 Flash** "
                               "in sidebar (most reliable format).")
                    with st.expander("🔍 Debug: see raw AI response (helps me fix the parser)"):
                        st.code(result[:3000])

    # ---- Questions display ----
    if "mock_q_list" in st.session_state and st.session_state["mock_q_list"]:
        q_list = st.session_state["mock_q_list"]
        meta = st.session_state.get("mock_meta", {})

        st.markdown("---")

        answered = sum(1 for i in range(len(q_list))
                       if st.session_state["q_state"].get(i, {}).get("feedback"))
        attempted = sum(1 for i in range(len(q_list))
                        if st.session_state["q_state"].get(i, {}).get("answer")) - answered
        pending = len(q_list) - answered - attempted

        st.markdown(
            f"### 📋 {len(q_list)} Questions "
            f"<span style='font-size:14px;color:#888'>"
            f"({meta.get('domain','')} • {meta.get('level','')})</span><br>"
            f"<span style='font-size:14px;font-weight:600;'>"
            f"<span style='color:#4CAF50'>✅ {answered} evaluated</span> • "
            f"<span style='color:#FF9800'>📝 {attempted} attempted</span> • "
            f"<span style='color:#9E9E9E'>⭕ {pending} pending</span></span>",
            unsafe_allow_html=True)

        if st.button("💾 Save full set to My Library", use_container_width=True, key="lib_save_q"):
            lib = st.session_state.setdefault("library", [])
            lib.append({"type": "questions", "meta": meta,
                        "content": st.session_state["mock_questions"]})
            st.success(f"✅ Saved! Library has {len(lib)} item(s).")

        # ---- Question dropdown ----
        st.markdown("#### 🎯 Pick Question to Practice")

        def _q_label(i):
            q = q_list[i]
            s = st.session_state["q_state"].get(i, {})
            if s.get("feedback"):
                icon = "✅"
            elif s.get("answer"):
                icon = "📝"
            else:
                icon = "⭕"
            return f"{icon} Q{i+1}: {q[:70]}{'...' if len(q) > 70 else ''}"

        current_idx = st.session_state.get("mock_q_idx", 0)

        # IMPORTANT: Use index= for initial value, capture return for change detection
        picked_idx = st.selectbox(
            "Pick question",
            options=list(range(len(q_list))),
            format_func=_q_label,
            index=current_idx,
            key=f"mock_q_picker_{current_idx}",  # dynamic key to force refresh
            label_visibility="collapsed"
        )
        if picked_idx != current_idx:
            _switch_to(picked_idx)
            st.rerun()

        # Prev / Counter / Next
        nav_c1, nav_c2, nav_c3 = st.columns([1, 2, 1])
        with nav_c1:
            if st.button("⬅️ Previous", use_container_width=True, key="prev_q_btn",
                         disabled=(current_idx == 0)):
                _switch_to(current_idx - 1)
                st.rerun()
        with nav_c2:
            st.markdown(
                f"<div style='text-align:center;padding:8px;background:#e8eaf6;"
                f"border-radius:6px;font-weight:600;font-size:15px;'>"
                f"Question {current_idx + 1} of {len(q_list)}</div>",
                unsafe_allow_html=True)
        with nav_c3:
            if st.button("Next ➡️", use_container_width=True, key="next_q_btn",
                         disabled=(current_idx >= len(q_list) - 1)):
                _switch_to(current_idx + 1)
                st.rerun()

        # ---- Active question card ----
        active_q = q_list[current_idx]
        st.markdown(
            f"<div style='background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);"
            f"padding:20px;border-radius:10px;margin-top:12px;color:white;'>"
            f"<div style='color:#FFEB3B;font-size:13px;font-weight:700;margin-bottom:6px;'>"
            f"❓ QUESTION {current_idx + 1}:</div>"
            f"<div style='font-size:16px;line-height:1.6;'>{active_q}</div></div>",
            unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 🎤 Answer with Voice or Type")

        st.info(
            "🍎 **iPhone Safari**: Best with **English (India)** — speak Hinglish naturally. "
            "**Android Chrome** supports Kannada/Tamil/Telugu directly. "
            "✨ **AI Polish** converts ANY mixed speech to professional English!"
        )

        vcol1, vcol2 = st.columns([3, 2])
        with vcol1:
            voice_lang = st.selectbox(
                "🌐 Speech language",
                ["en-IN — English (India) [RECOMMENDED]",
                 "en-US — English (US)",
                 "en-GB — English (UK)",
                 "en-AU — English (Australia)",
                 "hi-IN — Hindi",
                 "kn-IN — Kannada (Android Chrome best)",
                 "ta-IN — Tamil (Android Chrome best)",
                 "te-IN — Telugu (Android Chrome best)",
                 "ml-IN — Malayalam",
                 "mr-IN — Marathi",
                 "bn-IN — Bengali",
                 "gu-IN — Gujarati",
                 "pa-IN — Punjabi",
                 "ur-IN — Urdu",
                 "es-ES — Spanish",
                 "fr-FR — French",
                 "de-DE — German",
                 "ja-JP — Japanese",
                 "zh-CN — Chinese (Mandarin)",
                 "ar-SA — Arabic"],
                key="voice_lang_pick")
        with vcol2:
            st.markdown("<div style='padding-top:28px'></div>", unsafe_allow_html=True)
            st.caption("💡 Use mobile data, not office Wi-Fi")

        lang_code = voice_lang.split(" — ")[0]

        # ---- Voice recorder (just controls, no inner transcript) ----
        components.html(
            f"""
            <div style="font-family:system-ui,sans-serif;padding:14px;
                        background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);
                        border-radius:12px;color:white;">
                <button id="vbtn" style="
                    width:100%;padding:18px;font-size:17px;font-weight:700;
                    background:linear-gradient(135deg,#4CAF50 0%,#2E7D32 100%);
                    color:white;border:none;border-radius:10px;cursor:pointer;
                    box-shadow:0 4px 12px rgba(76,175,80,0.4);">
                    🎤 Tap to Start Recording
                </button>
                <div id="vstat" style="margin-top:10px;font-size:14px;color:#b0c4de;
                                      text-align:center;min-height:22px;font-weight:500;">
                    Tap to record • Transcript fills box below
                </div>
                <div style="margin-top:8px;display:flex;justify-content:space-between;
                            font-size:12px;color:#9aa5b8;">
                    <span id="vwc">📝 0 words</span>
                    <span id="vdur">⏱️ 0:00</span>
                </div>
            </div>
            <script>
            (function() {{
                const btn = document.getElementById('vbtn');
                const stat = document.getElementById('vstat');
                const wcEl = document.getElementById('vwc');
                const durEl = document.getElementById('vdur');
                const SR = window.SpeechRecognition || window.webkitSpeechRecognition;
                if (!SR) {{
                    btn.disabled = true;
                    btn.style.background = '#888';
                    btn.textContent = '❌ Voice not supported';
                    stat.textContent = 'Try Chrome (Android) or Safari (iPhone)';
                    return;
                }}
                const rec = new SR();
                rec.lang = '{lang_code}';
                rec.continuous = true;
                rec.interimResults = true;
                let finalText = '';
                let recording = false;
                let startTime = 0;
                let timerInterval = null;

                function findTranscriptBox() {{
                    try {{
                        const pDoc = window.parent.document;
                        const tas = pDoc.querySelectorAll('textarea[aria-label="Recording transcript (auto-filled, editable)"]');
                        return tas.length > 0 ? tas[0] : null;
                    }} catch(e) {{ return null; }}
                }}

                function fillBox(text) {{
                    const ta = findTranscriptBox();
                    if (!ta) return false;
                    try {{
                        const setter = Object.getOwnPropertyDescriptor(
                            window.parent.HTMLTextAreaElement.prototype, 'value'
                        ).set;
                        setter.call(ta, text);
                        ta.dispatchEvent(new Event('input', {{ bubbles: true }}));
                        return true;
                    }} catch(e) {{ return false; }}
                }}

                function doCopy(text) {{
                    if (navigator.clipboard && navigator.clipboard.writeText) {{
                        navigator.clipboard.writeText(text).catch(() => fallbackCopy(text));
                    }} else {{
                        fallbackCopy(text);
                    }}
                }}
                function fallbackCopy(text) {{
                    const ta = document.createElement('textarea');
                    ta.value = text;
                    ta.style.position = 'fixed'; ta.style.top = '-9999px';
                    document.body.appendChild(ta);
                    ta.select();
                    try {{ document.execCommand('copy'); }} catch(e) {{}}
                    document.body.removeChild(ta);
                }}

                function updateStats() {{
                    const words = finalText.trim().split(/\\s+/).filter(w => w.length > 0).length;
                    wcEl.textContent = '📝 ' + words + ' words';
                    if (recording && startTime) {{
                        const sec = Math.floor((Date.now() - startTime) / 1000);
                        const m = Math.floor(sec / 60);
                        const s = sec % 60;
                        durEl.textContent = '⏱️ ' + m + ':' + (s < 10 ? '0' : '') + s;
                    }}
                }}

                btn.onclick = () => {{
                    if (!recording) {{
                        finalText = '';
                        try {{
                            rec.start();
                            recording = true;
                            startTime = Date.now();
                            btn.textContent = '⏹️ Stop Recording';
                            btn.style.background = 'linear-gradient(135deg,#F44336 0%,#C62828 100%)';
                            stat.innerHTML = '<span style="color:#F44336;font-size:18px">●</span> ' +
                                             '<span style="color:#fff">Recording... speak now</span>';
                            timerInterval = setInterval(updateStats, 500);
                        }} catch(e) {{
                            stat.textContent = '⚠️ ' + e.message;
                        }}
                    }} else {{
                        rec.stop();
                        recording = false;
                        clearInterval(timerInterval);
                        btn.textContent = '🎤 Tap to Start Recording';
                        btn.style.background = 'linear-gradient(135deg,#4CAF50 0%,#2E7D32 100%)';
                        if (finalText.trim()) {{
                            doCopy(finalText.trim());
                            const filled = fillBox(finalText.trim());
                            stat.textContent = filled ? '✅ Filled below — tap Polish!' : '✅ Copied — paste manually below';
                        }} else {{
                            stat.textContent = '⚠️ No speech detected';
                        }}
                        updateStats();
                    }}
                }};
                rec.onresult = (event) => {{
                    let interim = '';
                    for (let i = event.resultIndex; i < event.results.length; i++) {{
                        const t = event.results[i][0].transcript;
                        if (event.results[i].isFinal) finalText += t + ' ';
                        else interim += t;
                    }}
                    // Live update the box as user speaks
                    fillBox(finalText + interim);
                    updateStats();
                }};
                rec.onerror = (e) => {{
                    stat.textContent = '⚠️ Error: ' + e.error + ' — allow mic';
                    recording = false;
                    clearInterval(timerInterval);
                    btn.textContent = '🎤 Tap to Start Recording';
                    btn.style.background = 'linear-gradient(135deg,#4CAF50 0%,#2E7D32 100%)';
                }};
                rec.onend = () => {{
                    if (recording) {{
                        try {{ rec.start(); }} catch(e) {{}}
                    }}
                }};
            }})();
            </script>
            """,
            height=200,
        )

        # ---- FORM wrap: forces Streamlit to read CURRENT textarea value ----
        # This fixes "Transcript box is empty" bug — st.form re-reads the DOM
        # value when form_submit_button is clicked, even after JS injection.
        with st.form("voice_polish_form", clear_on_submit=False):
            st.text_area(
                "Recording transcript (auto-filled, editable)",
                height=180, key="raw_voice_in",
                placeholder="Your voice transcript appears here LIVE while you speak. "
                            "Edit if needed, then tap Polish or Use Raw."
            )
            st.caption("👆 If text isn't here after speaking, long-press box → Paste (clipboard backup).")

            pcol1, pcol2 = st.columns(2)
            with pcol1:
                polish_clicked = st.form_submit_button(
                    "✨ Polish My Speech", use_container_width=True, type="primary")
            with pcol2:
                use_raw_clicked = st.form_submit_button(
                    "⚡ Use Raw Speech", use_container_width=True)

        # After form submit, Streamlit has synced the DOM value to session_state
        current_transcript = st.session_state.get("raw_voice_in", "").strip()

        if polish_clicked:   
            if not current_transcript:
                st.warning("⚠️ Transcript box is empty. Speak or type first.")
            else:
                polish_prompt = (
                    "You are a professional interview answer editor. "
                    "User has spoken an answer that may mix languages "
                    "(English + Hindi/Kannada/Tamil/Telugu), with grammar mistakes, "
                    "filler words, and rough phrasing.\n\n"
                    "Convert into polished professional English. Use STAR if possible. "
                    "Remove fillers. Fix grammar. Translate non-English to English. "
                    "Keep 3-6 sentences. Do NOT fabricate. Output ONLY polished answer.\n\n"
                    "INTERVIEW QUESTION: " + active_q + "\n\n"
                    "---ROUGH TRANSCRIPT---\n" + current_transcript +
                    "\n\n---POLISHED ANSWER---"
                )
                with st.spinner("✨ AI polishing your answer..."):
                    polished = call_ai(polish_prompt, temperature=0.4)
                if polished:
                    st.session_state["polished_preview"] = polished.strip()
                    st.rerun()

        if use_raw_clicked:
            if not current_transcript:
                st.warning("⚠️ Transcript box is empty. Speak or type first.")
            else:
                st.session_state["mock_a_in"] = current_transcript
                st.session_state.pop("polished_preview", None)
                _save_current_q()
                st.success("✅ Raw transcript set as answer.")
                st.rerun()

        # ---- Polished preview ----
        if "polished_preview" in st.session_state:
            st.markdown("---")
            st.markdown("##### ✨ Polished Version (review then use)")
            st.markdown(
                f"<div style='background:#E8F5E9;padding:14px;border-radius:8px;"
                f"border-left:4px solid #4CAF50;color:#1a1a2e;font-size:15px;line-height:1.6;'>"
                f"{st.session_state['polished_preview']}</div>",
                unsafe_allow_html=True)
            pp1, pp2, pp3 = st.columns(3)
            with pp1:
                if st.button("✅ Use This Polished", use_container_width=True,
                             type="primary", key="use_polished_btn"):
                    st.session_state["mock_a_in"] = st.session_state["polished_preview"]
                    st.session_state.pop("polished_preview", None)
                    _save_current_q()
                    st.rerun()
            with pp2:
                if st.button("✏️ Edit First", use_container_width=True, key="edit_polished_btn"):
                    st.session_state["mock_a_in"] = st.session_state["polished_preview"]
                    st.session_state.pop("polished_preview", None)
                    _save_current_q()
                    st.rerun()
            with pp3:
                if st.button("🗑️ Discard", use_container_width=True, key="discard_polished_btn"):
                    st.session_state.pop("polished_preview", None)
                    st.rerun()

        st.markdown("---")
        st.markdown("##### 📝 Your Answer (auto-filled, editable)")
        st.text_area("Your Answer", height=200, key="mock_a_in",
                     label_visibility="collapsed")

        if st.button("📊 Evaluate My Answer", type="primary",
                     use_container_width=True, key="eval_btn"):
            current_answer = st.session_state.get("mock_a_in", "").strip()
            if not current_answer:
                st.warning("Please provide an answer first.")
            else:
                result = call_ai(p_mock_eval(active_q, current_answer, exp_level))
                if result:
                    st.session_state["mock_feedback"] = result
                    _save_current_q()

        if "mock_feedback" in st.session_state:
            st.markdown("### 📊 AI Feedback")
            st.text_area("Feedback", st.session_state["mock_feedback"], height=400,
                         key="mock_fb_out", label_visibility="collapsed")
            download_buttons(st.session_state["mock_feedback"], region,
                             "Mock_Feedback", name_override="Candidate")

            fb_c1, fb_c2, fb_c3 = st.columns(3)
            with fb_c1:
                if st.button("💾 Save to Library", use_container_width=True, key="lib_save_fb"):
                    lib = st.session_state.setdefault("library", [])
                    lib.append({
                        "type": "feedback",
                        "meta": {**(st.session_state.get("mock_meta") or {}),
                                 "question": active_q},
                        "content": st.session_state["mock_feedback"],
                    })
                    st.success(f"✅ Saved! Library has {len(lib)} item(s).")
            with fb_c2:
                next_disabled = current_idx >= len(q_list) - 1
                if st.button("🎯 Next Question →", use_container_width=True, type="primary",
                             key="next_q_eval_btn", disabled=next_disabled):
                    _switch_to(current_idx + 1)
                    st.rerun()
            with fb_c3:
                if st.button("🔁 Try Again", use_container_width=True, key="retry_btn"):
                    for k in ["raw_voice_in", "mock_a_in", "mock_feedback", "polished_preview"]:
                        st.session_state.pop(k, None)
                    st.session_state["q_state"].pop(current_idx, None)
                    st.rerun()

# ============================================================
# PAGE: COACHING
# ============================================================
elif page == "🧑‍💼 Coaching":
    st.title("🧑‍💼 Career Coaching")
    topic = st.selectbox("Topic",
        ["Career Change", "Salary Negotiation", "Skill Development", "Job Search",
         "Interview Confidence", "LinkedIn", "Networking", "Promotion",
         "Work-Life Balance", "Remote Work"], key="coach_topic")
    context = st.text_area("Your situation/context", height=200, key="coach_ctx",
                           placeholder="Describe where you are, what you want, and challenges...")
    if st.button("💡 Get Advice", type="primary", use_container_width=True):
        if not context.strip():
            st.warning("Please describe your situation.")
        else:
            result = call_ai(p_coach(topic, context))
            if result:
                st.session_state["coach_result"] = result
    if "coach_result" in st.session_state:
        st.markdown("### 💡 Personalised Advice")
        st.text_area("Advice", st.session_state["coach_result"], height=500,
                     key="coach_out", label_visibility="collapsed")
        download_buttons(st.session_state["coach_result"], region,
                         "Career_Advice", name_override="Candidate")

# ============================================================
# PAGE: MY LIBRARY
# ============================================================
elif page == "📚 My Library":
    st.title("📚 My Library")
    st.caption("Save your mock interview questions & feedback. Export as JSON.")

    st.info("💡 **How it works:** Items are saved during your current session. "
            "Before closing, click **⬇️ Export Library** to download a backup. "
            "Next time, **⬆️ Import** to restore.")

    lib = st.session_state.setdefault("library", [])
    st.success(f"📚 Your library has **{len(lib)}** saved item(s).")

    col1, col2 = st.columns(2)
    with col1:
        if lib:
            st.download_button("⬇️ Export Library (JSON backup)",
                               data=json.dumps(lib, indent=2, ensure_ascii=False),
                               file_name=f"my_library_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                               mime="application/json", use_container_width=True)
        else:
            st.button("⬇️ Export (nothing yet)", disabled=True, use_container_width=True)
    with col2:
        if lib and st.button("🗑️ Clear ALL items", use_container_width=True):
            st.session_state["library"] = []
            st.rerun()

    st.markdown("**⬆️ Import previously exported library:**")
    up = st.file_uploader("Upload library JSON", type=["json"], key="lib_import")
    if up is not None:
        try:
            data = json.loads(up.read().decode("utf-8"))
            if isinstance(data, list):
                st.session_state["library"] = data + lib
                st.success(f"✅ Imported {len(data)} item(s).")
                st.rerun()
            else:
                st.error("Bad file: expected a JSON list.")
        except Exception as e:
            st.error(f"Bad file: {e}")

    st.markdown("---")
    if not lib:
        st.info("📭 Library is empty. Go to **🎙️ Mock Interview** to save items.")
    else:
        st.markdown("### 📑 Your saved items")
        for i, item in enumerate(reversed(lib)):
            real_idx = len(lib) - 1 - i
            meta = item.get("meta", {})
            with st.expander(f"#{real_idx+1} • {item.get('type','?').title()} • "
                             f"{meta.get('domain','?')} • {meta.get('level','?')} • "
                             f"{meta.get('ts','?')}"):
                st.text_area("Content", item.get("content", ""), height=300,
                             key=f"lib_view_{real_idx}", label_visibility="collapsed")
                cdl1, cdl2 = st.columns(2)
                with cdl1:
                    st.download_button("💾 Download .txt", item.get("content", ""),
                                       file_name=f"library_item_{real_idx+1}.txt",
                                       mime="text/plain", use_container_width=True,
                                       key=f"lib_dl_{real_idx}")
                with cdl2:
                    if st.button("🗑️ Delete", use_container_width=True, key=f"lib_del_{real_idx}"):
                        lib.pop(real_idx)
                        st.rerun()

# ============================================================
# PAGE: SETTINGS
# ============================================================
elif page == "⚙️ Settings":
    st.title("⚙️ Settings & Info")
    with st.container(border=True):
        st.markdown("### 🤖 AI Providers")
        st.markdown("- **Gemini (FREE):** https://aistudio.google.com/apikey")
        st.markdown("- **OpenAI (Paid):** https://platform.openai.com/api-keys")
        st.markdown("- **Claude / Anthropic (Paid):** https://console.anthropic.com/settings/keys")
    with st.container(border=True):
        st.markdown("### 🌍 Region-Aware CV Generation")
        st.markdown("- **UK / US / Canada / Australia:** NO photo")
        st.markdown("- **Germany / Japan / China / Spain / Italy / UAE:** Photo placeholder")
        st.markdown("- **India:** Declaration at bottom")
        st.markdown("- **Italy:** GDPR clause")
        st.markdown("- **France / EU:** CEFR language levels")
        st.markdown("- **UAE / Gulf:** Nationality + visa status mandatory")
    with st.container(border=True):
        st.markdown("### 🎙️ Mock Interview Features")
        st.markdown(f"**{sum(len(v) for v in DOMAIN_GROUPS.values())}** "
                    f"domains across **{len(DOMAIN_GROUPS)}** categories.")
        st.markdown("🎤 **Voice input** in 20+ languages (Eng/Hindi/Kannada/Tamil/Telugu/+).")
        st.markdown("✨ **AI Polish** converts rough/mixed speech to professional English.")
        st.markdown("🔄 **Anti-repeat**: each generated question tracked + AI explicitly excluded "
                    "from future rounds. Tap 🧹 Clear to reset.")
        st.markdown("📊 History counter shows tracked questions in real time.")
    with st.container(border=True):
        st.markdown("### 📚 My Library")
        st.markdown("Save mock interview Qs & feedback during your session. "
                    "**Always export to JSON before closing** — items reset on leave.")
    with st.container(border=True):
        st.markdown("### 📱 Use on Phone")
        st.markdown("**iPhone (Safari):** Share → Add to Home Screen")
        st.markdown("**Android (Chrome):** Menu → Add to Home Screen")
        st.markdown("Tip: tap any sidebar tab — it auto-closes on mobile.")
    docx_status = "OK" if HAS_DOCX else "Missing"
    rl_status   = "OK" if HAS_RL   else "Missing"
    pdf_status  = "OK" if HAS_PDF  else "Missing"
    req_status  = "OK" if HAS_REQ  else "Missing"
    st.caption(f"SSL: {SSL_M} | python-docx: {docx_status} | reportlab: {rl_status} "
               f"| pypdf: {pdf_status} | requests: {req_status}")
